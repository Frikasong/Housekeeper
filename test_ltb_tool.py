#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Comprehensive test suite for the LTB Evidence Brief Generator.

Exercises upload, PDF generation, evidence analysis, and smoke-test routes.
Produces structured JSON and human-readable TXT logs in test_logs/.

Usage:
    python test_ltb_tool.py                 # fast tests only (~30s)
    python test_ltb_tool.py --include-slow  # include semantic model tests (~2min)
"""

import argparse
import datetime
import io
import json
import os
import platform
import shutil
import sys
import tempfile
import time
import traceback

# ---------------------------------------------------------------------------
# Ensure project root is importable
# ---------------------------------------------------------------------------
PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
if PROJECT_DIR not in sys.path:
    sys.path.insert(0, PROJECT_DIR)

from PIL import Image as PILImage, ImageDraw, ImageFont
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
    Image as RLImage,
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_RIGHT

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pillow_heif
    pillow_heif.register_heif_opener()
    HAS_PILLOW_HEIF = True
except ImportError:
    HAS_PILLOW_HEIF = False


# ═══════════════════════════════════════════════════════════════════════════════
# TEST FILE GENERATORS — realistic fake evidence files
# ═══════════════════════════════════════════════════════════════════════════════

def _draw_text(draw, text, xy, fill="black", font_size=16):
    """Draw text using default font (no TTF dependency)."""
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", font_size)
    except (IOError, OSError):
        font = ImageFont.load_default()
    draw.text(xy, text, fill=fill, font=font)


def make_png(width=100, height=100, mode="RGB"):
    """Create a fake photo of a rental unit hallway — shapes and labels."""
    if mode == "RGBA":
        img = PILImage.new("RGBA", (width, height), (245, 240, 230, 255))
    else:
        img = PILImage.new("RGB", (width, height), (245, 240, 230))
    draw = ImageDraw.Draw(img)
    # Floor
    draw.rectangle([0, int(height * 0.75), width, height], fill=(180, 160, 140))
    # Door
    dw, dh = int(width * 0.25), int(height * 0.45)
    dx = int(width * 0.15)
    dy = int(height * 0.30)
    draw.rectangle([dx, dy, dx + dw, dy + dh], fill=(120, 80, 50), outline=(80, 50, 30), width=2)
    draw.ellipse([dx + dw - 12, dy + dh // 2 - 4, dx + dw - 4, dy + dh // 2 + 4], fill="gold")
    # Window
    wx = int(width * 0.60)
    wy = int(height * 0.25)
    ww, wh = int(width * 0.25), int(height * 0.25)
    draw.rectangle([wx, wy, wx + ww, wy + wh], fill=(200, 220, 255), outline=(100, 100, 100), width=2)
    draw.line([wx + ww // 2, wy, wx + ww // 2, wy + wh], fill=(100, 100, 100), width=1)
    draw.line([wx, wy + wh // 2, wx + ww, wy + wh // 2], fill=(100, 100, 100), width=1)
    # Label
    if width >= 200:
        _draw_text(draw, "Unit 4B — Hallway", (10, 5), fill=(80, 80, 80), font_size=14)
        _draw_text(draw, "2026-01-15", (10, height - 22), fill=(120, 120, 120), font_size=12)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    buf.name = "unit_hallway_photo.png"
    return buf


def make_jpeg(width=800, height=600):
    """Create a fake photo of water damage in a kitchen ceiling."""
    img = PILImage.new("RGB", (width, height), (250, 248, 240))
    draw = ImageDraw.Draw(img)
    # Ceiling area
    draw.rectangle([0, 0, width, int(height * 0.5)], fill=(230, 225, 215))
    # Water stain — irregular brownish patch
    import random
    rng = random.Random(42)
    cx, cy = width // 2, int(height * 0.28)
    for _ in range(300):
        rx = cx + rng.randint(-120, 120)
        ry = cy + rng.randint(-60, 60)
        r = rng.randint(3, 12)
        shade = rng.randint(140, 190)
        draw.ellipse([rx - r, ry - r, rx + r, ry + r], fill=(shade, shade - 20, shade - 50))
    # Cabinets
    for i in range(3):
        x0 = int(width * 0.1) + i * int(width * 0.28)
        y0 = int(height * 0.55)
        cw, ch = int(width * 0.22), int(height * 0.35)
        draw.rectangle([x0, y0, x0 + cw, y0 + ch], fill=(160, 130, 90), outline=(100, 70, 40), width=2)
        draw.rectangle([x0 + 4, y0 + 4, x0 + cw - 4, y0 + ch - 4], outline=(130, 100, 60), width=1)
    # Labels
    _draw_text(draw, "Kitchen ceiling — water damage", (15, 10), fill=(80, 0, 0), font_size=18)
    _draw_text(draw, "123 Test St, Unit 4B", (15, 35), fill=(100, 100, 100), font_size=14)
    _draw_text(draw, "Photo taken: 2026-02-10", (15, height - 28), fill=(120, 120, 120), font_size=13)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=88)
    buf.seek(0)
    buf.name = "kitchen_water_damage.jpg"
    return buf


def make_webp(width=100, height=100):
    """Create a fake photo of a broken window latch."""
    img = PILImage.new("RGB", (width, height), (220, 225, 230))
    draw = ImageDraw.Draw(img)
    # Window frame
    m = max(4, width // 20)
    draw.rectangle([m, m, width - m, height - m], outline=(90, 90, 90), width=3)
    draw.line([width // 2, m, width // 2, height - m], fill=(90, 90, 90), width=2)
    draw.line([m, height // 2, width - m, height // 2], fill=(90, 90, 90), width=2)
    # Broken latch — red X
    lx, ly = int(width * 0.5), int(height * 0.5)
    s = max(6, width // 12)
    draw.line([lx - s, ly - s, lx + s, ly + s], fill="red", width=3)
    draw.line([lx - s, ly + s, lx + s, ly - s], fill="red", width=3)
    buf = io.BytesIO()
    img.save(buf, format="WEBP")
    buf.seek(0)
    buf.name = "broken_window_latch.webp"
    return buf


def make_pdf_text(text=None):
    """Create a fake rent payment record PDF with a table of monthly payments."""
    if text and text != "This is a test PDF document for LTB evidence.":
        # Caller wants custom text — simple single-paragraph PDF
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter)
        styles = getSampleStyleSheet()
        doc.build([Paragraph(text, styles["Normal"])])
        buf.seek(0)
        buf.name = "document.pdf"
        return buf

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
                            topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    title_s = ParagraphStyle("RRTitle", parent=styles["Title"], fontSize=16,
                             spaceAfter=6)
    normal_s = styles["Normal"]
    story = []

    story.append(Paragraph("Rent Payment Record", title_s))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.black))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph("Tenant: Jane Doe", normal_s))
    story.append(Paragraph("Unit: 123 Test Street, Unit 4B, Toronto ON M5V 2T6", normal_s))
    story.append(Paragraph("Landlord: ABC Property Management Inc.", normal_s))
    story.append(Paragraph("Monthly Rent: $1,850.00", normal_s))
    story.append(Spacer(1, 0.2 * inch))

    data = [["Month", "Due Date", "Amount Due", "Amount Paid", "Date Paid", "Balance"]]
    records = [
        ("Jan 2026", "Jan 1", "$1,850.00", "$1,850.00", "Jan 1", "$0.00"),
        ("Feb 2026", "Feb 1", "$1,850.00", "$1,850.00", "Feb 3", "$0.00"),
        ("Mar 2026", "Mar 1", "$1,850.00", "$900.00",   "Mar 5", "$950.00"),
        ("Apr 2026", "Apr 1", "$1,850.00", "$0.00",     "—",     "$2,800.00"),
        ("May 2026", "May 1", "$1,850.00", "$0.00",     "—",     "$4,650.00"),
    ]
    for row in records:
        data.append(list(row))

    col_w = [1.0 * inch, 0.85 * inch, 1.0 * inch, 1.0 * inch, 0.85 * inch, 1.0 * inch]
    tbl = Table(data, colWidths=col_w)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#003366")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
        ("ALIGN", (2, 1), (-1, -1), "RIGHT"),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph(
        "Total arrears as of May 31, 2026: <b>$4,650.00</b>",
        ParagraphStyle("Arrears", parent=normal_s, fontSize=11)))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph(
        "This record was prepared by ABC Property Management Inc. for filing "
        "with the Landlord and Tenant Board in support of application TSL-99999-26.",
        ParagraphStyle("Note", parent=normal_s, fontSize=9, textColor=colors.grey)))

    doc.build(story)
    buf.seek(0)
    buf.name = "rent_payment_record.pdf"
    return buf


def make_pdf_with_image():
    """Create a fake maintenance inspection report PDF with a photo."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
                            topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Maintenance Inspection Report", styles["Title"]))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.black))
    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph("Property: 123 Test Street, Unit 4B, Toronto ON", styles["Normal"]))
    story.append(Paragraph("Inspector: Mike Chen, Building Superintendent", styles["Normal"]))
    story.append(Paragraph("Date of Inspection: February 12, 2026", styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(
        "During the scheduled maintenance inspection, the following issues were "
        "identified in the tenant's unit:", styles["Normal"]))
    story.append(Spacer(1, 0.1 * inch))
    story.append(Paragraph("1. Water stain on kitchen ceiling (approx. 18 inches diameter)", styles["Normal"]))
    story.append(Paragraph("2. Window latch broken in bedroom — unit not secure", styles["Normal"]))
    story.append(Paragraph("3. Bathroom exhaust fan inoperative", styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))

    # Embed a small photo placeholder
    photo = PILImage.new("RGB", (300, 200), (230, 220, 210))
    d = ImageDraw.Draw(photo)
    d.rectangle([20, 20, 280, 180], outline=(150, 130, 110), width=2)
    _draw_text(d, "Ceiling damage photo", (50, 80), fill=(100, 70, 40), font_size=16)
    img_buf = io.BytesIO()
    photo.save(img_buf, "JPEG", quality=85)
    img_buf.seek(0)
    story.append(RLImage(img_buf, width=3 * inch, height=2 * inch))

    story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph(
        "Recommended action: Repair ceiling leak from unit above (5B), replace "
        "window latch hardware, service exhaust fan motor.", styles["Normal"]))

    doc.build(story)
    buf.seek(0)
    buf.name = "maintenance_inspection_report.pdf"
    return buf


def make_docx(text=None):
    """Create a fake N4 notice or demand letter as a DOCX file."""
    if not HAS_DOCX:
        return None
    doc = DocxDocument()

    if text and text != "This is a test DOCX document for LTB evidence.":
        doc.add_paragraph(text)
    else:
        doc.add_heading("DEMAND LETTER — Rent Arrears", level=1)
        doc.add_paragraph("")
        doc.add_paragraph("Date: March 1, 2026")
        doc.add_paragraph("To: Jane Doe")
        doc.add_paragraph("Re: Rent Arrears — 123 Test Street, Unit 4B, Toronto ON M5V 2T6")
        doc.add_paragraph("")
        doc.add_paragraph(
            "Dear Ms. Doe,"
        )
        doc.add_paragraph(
            "This letter serves as formal notice that your rent account is in arrears. "
            "As of February 28, 2026, the total outstanding balance is $2,800.00, "
            "representing unpaid rent for March and a partial payment shortfall from "
            "the preceding month."
        )
        doc.add_paragraph(
            "Under the Residential Tenancies Act, 2006, you are required to pay rent "
            "on the first day of each month. Your lease agreement dated September 1, "
            "2024 specifies a monthly rent of $1,850.00."
        )
        doc.add_paragraph(
            "Please remit the full outstanding amount within 14 days of this letter "
            "to avoid further legal action, which may include an application to the "
            "Landlord and Tenant Board for termination of the tenancy."
        )
        doc.add_paragraph("")
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph("Robert Smith")
        doc.add_paragraph("ABC Property Management Inc.")
        doc.add_paragraph("Tel: (416) 555-0199")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    buf.name = "demand_letter_rent_arrears.docx"
    return buf


def make_md(content=None):
    """Create a fake tenant communication log as a Markdown file."""
    if content is None:
        content = (
            "# Tenant Communication Log\n\n"
            "**Unit:** 123 Test Street, Unit 4B, Toronto ON\n"
            "**Tenant:** Jane Doe\n"
            "**File:** TSL-99999-26\n\n"
            "---\n\n"
            "## January 15, 2026\n"
            "- Tenant emailed requesting repair of kitchen ceiling leak.\n"
            "- Superintendent dispatched same day; identified source as unit 5B above.\n\n"
            "## February 3, 2026\n"
            "- Rent received 2 days late ($1,850.00 e-transfer).\n"
            "- Landlord sent reminder email about due-date policy.\n\n"
            "## March 1, 2026\n"
            "- Partial rent received: $900.00 of $1,850.00.\n"
            "- Landlord called tenant; tenant stated financial hardship.\n\n"
            "## March 5, 2026\n"
            "- N4 Notice of Termination served for non-payment of rent.\n"
            "- Copy sent via registered mail and email.\n\n"
            "## March 10, 2026\n"
            "- Tenant acknowledged receipt of N4; requested payment plan.\n"
            "- Landlord declined; filed application with LTB.\n"
        )
    buf = io.BytesIO(content.encode("utf-8"))
    buf.name = "tenant_communication_log.md"
    return buf


def make_gif(width=400, height=300):
    """Create a fake GIF photo of a damaged front door."""
    img = PILImage.new("RGB", (width, height), (210, 200, 185))
    draw = ImageDraw.Draw(img)
    # Door frame
    dx, dw = int(width * 0.25), int(width * 0.5)
    dy, dh = int(height * 0.1), int(height * 0.85)
    draw.rectangle([dx, dy, dx + dw, dy + dh], fill=(130, 85, 55), outline=(80, 50, 30), width=3)
    # Damage scratches — red diagonal lines
    for i in range(5):
        x0 = dx + 30 + i * 25
        y0 = dy + 60 + i * 15
        draw.line([x0, y0, x0 + 40, y0 + 60], fill=(180, 30, 30), width=3)
    # Knob
    draw.ellipse([dx + dw - 45, dy + dh // 2 - 8, dx + dw - 25, dy + dh // 2 + 8], fill="gold")
    if width >= 200:
        _draw_text(draw, "Front door — scratch damage", (10, 5), fill=(80, 0, 0), font_size=13)
        _draw_text(draw, "2026-02-20", (10, height - 22), fill=(120, 120, 120), font_size=11)
    buf = io.BytesIO()
    img.save(buf, format="GIF")
    buf.seek(0)
    buf.name = "front_door_damage.gif"
    return buf


def make_bmp(width=400, height=300):
    """Create a fake BMP photo of a cracked bathroom tile."""
    img = PILImage.new("RGB", (width, height), (235, 235, 230))
    draw = ImageDraw.Draw(img)
    # Tile grid
    tile_size = max(40, width // 8)
    for x in range(0, width, tile_size):
        draw.line([x, 0, x, height], fill=(200, 200, 195), width=1)
    for y in range(0, height, tile_size):
        draw.line([0, y, width, y], fill=(200, 200, 195), width=1)
    # Crack pattern — jagged dark line through center
    import random
    rng = random.Random(99)
    cx, cy = width // 3, height // 4
    points = [(cx, cy)]
    for _ in range(12):
        cx += rng.randint(10, 30)
        cy += rng.randint(-10, 25)
        points.append((cx, cy))
    draw.line(points, fill=(60, 40, 30), width=2)
    # Branch cracks
    for px, py in points[2::3]:
        bx, by = px + rng.randint(-20, 20), py + rng.randint(10, 30)
        draw.line([px, py, bx, by], fill=(90, 70, 50), width=1)
    if width >= 200:
        _draw_text(draw, "Bathroom tile — cracked", (10, 5), fill=(80, 0, 0), font_size=13)
        _draw_text(draw, "2026-01-28", (10, height - 22), fill=(120, 120, 120), font_size=11)
    buf = io.BytesIO()
    img.save(buf, format="BMP")
    buf.seek(0)
    buf.name = "bathroom_tile_cracked.bmp"
    return buf


def make_tiff(width=400, height=300):
    """Create a fake TIFF photo of mold on a bedroom wall."""
    img = PILImage.new("RGB", (width, height), (245, 242, 235))
    draw = ImageDraw.Draw(img)
    # Wall with baseboard
    draw.rectangle([0, int(height * 0.88), width, height], fill=(160, 140, 120))
    # Mold patches — dark greenish-brown splotches
    import random
    rng = random.Random(77)
    for _ in range(20):
        mx = rng.randint(int(width * 0.2), int(width * 0.8))
        my = rng.randint(int(height * 0.15), int(height * 0.7))
        r = rng.randint(8, 25)
        shade_r = rng.randint(40, 80)
        shade_g = rng.randint(50, 90)
        draw.ellipse([mx - r, my - r, mx + r, my + r], fill=(shade_r, shade_g, 30))
    if width >= 200:
        _draw_text(draw, "Bedroom wall — mold growth", (10, 5), fill=(80, 0, 0), font_size=13)
        _draw_text(draw, "2026-03-01", (10, height - 22), fill=(120, 120, 120), font_size=11)
    buf = io.BytesIO()
    img.save(buf, format="TIFF")
    buf.seek(0)
    buf.name = "bedroom_mold_growth.tiff"
    return buf


def make_heic(width=400, height=300):
    """Create a fake HEIC photo of a stained carpet. Requires pillow-heif."""
    if not HAS_PILLOW_HEIF:
        return None
    img = PILImage.new("RGB", (width, height), (180, 170, 150))
    draw = ImageDraw.Draw(img)
    # Carpet texture — horizontal lines
    for y in range(0, height, 4):
        shade = 175 + (y % 8)
        draw.line([0, y, width, y], fill=(shade, shade - 10, shade - 20), width=1)
    # Large dark stain
    import random
    rng = random.Random(55)
    cx, cy = width // 2, height // 2
    for _ in range(200):
        rx = cx + rng.randint(-80, 80)
        ry = cy + rng.randint(-50, 50)
        r = rng.randint(3, 10)
        shade = rng.randint(80, 120)
        draw.ellipse([rx - r, ry - r, rx + r, ry + r], fill=(shade, shade - 10, shade - 30))
    if width >= 200:
        _draw_text(draw, "Living room — carpet stain", (10, 5), fill=(80, 0, 0), font_size=13)
        _draw_text(draw, "2026-02-14", (10, height - 22), fill=(120, 120, 120), font_size=11)
    buf = io.BytesIO()
    pillow_heif.from_pillow(img).save(buf, format="HEIF")
    buf.seek(0)
    buf.name = "carpet_stain.heic"
    return buf


# ═══════════════════════════════════════════════════════════════════════════════
# TEST LOGGER
# ═══════════════════════════════════════════════════════════════════════════════

class TestLogger:
    """Captures every test result for console, TXT, and JSON output."""

    def __init__(self, log_dir):
        self.log_dir = log_dir
        os.makedirs(log_dir, exist_ok=True)
        self.run_id = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        self.run_started = datetime.datetime.utcnow().isoformat() + "Z"
        self.tests = []
        self.counts = {"total": 0, "passed": 0, "failed": 0, "skipped": 0}

        # Console + TXT buffer
        self._txt_lines = []
        self._log(f"=== LTB Tool Test Run {self.run_id} ===")
        self._log(f"Python {platform.python_version()} on {platform.system()} {platform.release()}")
        self._log(f"Started: {self.run_started}")
        self._log("")

    def _log(self, msg):
        print(msg)
        self._txt_lines.append(msg)

    def record(self, test_id, category, name, description, action,
               status, detail="", files_generated=None, started=None, finished=None):
        duration_ms = 0
        if started and finished:
            duration_ms = round((finished - started) * 1000, 1)

        entry = {
            "id": test_id,
            "category": category,
            "name": name,
            "description": description,
            "action": action,
            "started": datetime.datetime.utcfromtimestamp(started).isoformat() + "Z" if started else "",
            "finished": datetime.datetime.utcfromtimestamp(finished).isoformat() + "Z" if finished else "",
            "duration_ms": duration_ms,
            "status": status,
            "detail": detail,
            "files_generated": files_generated or [],
        }
        self.tests.append(entry)
        self.counts["total"] += 1
        if status == "PASS":
            self.counts["passed"] += 1
            icon = "OK"
        elif status == "FAIL":
            self.counts["failed"] += 1
            icon = "FAIL"
        else:
            self.counts["skipped"] += 1
            icon = "SKIP"

        self._log(f"  [{icon:4s}] {test_id:8s} {name} ({duration_ms:.0f}ms)")
        if status == "FAIL" and detail:
            for line in detail.split("\n")[:5]:
                self._log(f"           {line}")

    def finish(self):
        run_finished = datetime.datetime.utcnow().isoformat() + "Z"
        t0 = datetime.datetime.fromisoformat(self.run_started.replace("Z", "+00:00"))
        t1 = datetime.datetime.fromisoformat(run_finished.replace("Z", "+00:00"))
        duration = (t1 - t0).total_seconds()

        self._log("")
        self._log("=" * 60)
        s = self.counts
        self._log(
            f"TOTAL: {s['total']}  |  PASSED: {s['passed']}  |  "
            f"FAILED: {s['failed']}  |  SKIPPED: {s['skipped']}"
        )
        self._log(f"Duration: {duration:.1f}s")
        self._log("=" * 60)

        # Write JSON log
        json_path = os.path.join(self.log_dir, f"test_run_{self.run_id}.json")
        json_data = {
            "run_id": self.run_id,
            "run_started": self.run_started,
            "run_finished": run_finished,
            "duration_seconds": round(duration, 2),
            "python_version": platform.python_version(),
            "summary": self.counts,
            "tests": self.tests,
        }
        with open(json_path, "w") as f:
            json.dump(json_data, f, indent=2)

        # Write TXT log
        txt_path = os.path.join(self.log_dir, f"test_run_{self.run_id}.txt")
        with open(txt_path, "w") as f:
            f.write("\n".join(self._txt_lines) + "\n")

        self._log(f"\nLogs written to:")
        self._log(f"  {json_path}")
        self._log(f"  {txt_path}")

        return self.counts["failed"]


# ═══════════════════════════════════════════════════════════════════════════════
# LTB TEST FIXTURE
# ═══════════════════════════════════════════════════════════════════════════════

class LTBTestFixture:
    """Sets up Flask test client with an isolated temp upload directory."""

    def __init__(self):
        self.tmp_dir = None
        self.client = None
        self.app = None

    def setup(self):
        from app import app
        self.app = app
        self.tmp_dir = tempfile.mkdtemp(prefix="ltb_test_")
        self.app.config["UPLOAD_FOLDER"] = self.tmp_dir
        self.app.config["TESTING"] = True
        self.client = self.app.test_client()
        return self

    def teardown(self):
        if self.tmp_dir and os.path.exists(self.tmp_dir):
            shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def upload_file(self, file_buf, filename=None):
        """Helper: upload a single file and return parsed JSON response."""
        if filename is None:
            filename = getattr(file_buf, "name", "file.bin")
        file_buf.seek(0)
        resp = self.client.post(
            "/upload",
            data={"files": (file_buf, filename)},
            content_type="multipart/form-data",
        )
        return resp

    def upload_files_multi(self, file_list):
        """Upload multiple files in one POST. file_list = [(buf, name), ...]"""
        data = {}
        files = []
        for buf, name in file_list:
            buf.seek(0)
            files.append((buf, name))
        resp = self.client.post(
            "/upload",
            data={"files": [(buf, name) for buf, name in files]},
            content_type="multipart/form-data",
        )
        return resp

    def generate_brief(self, case_info, tabs):
        """Helper: call /generate with given case info and tabs."""
        resp = self.client.post(
            "/generate",
            data=json.dumps({"case_info": case_info, "tabs": tabs}),
            content_type="application/json",
        )
        return resp

    def list_upload_files(self):
        """List files in the temp upload directory."""
        if self.tmp_dir and os.path.exists(self.tmp_dir):
            return os.listdir(self.tmp_dir)
        return []


# ═══════════════════════════════════════════════════════════════════════════════
# TEST RUNNER
# ═══════════════════════════════════════════════════════════════════════════════

class TestRunner:
    """Runs all tests in sequence, catching failures without aborting."""

    def __init__(self, include_slow=False):
        self.include_slow = include_slow
        self.log_dir = os.path.join(PROJECT_DIR, "test_logs")
        self.logger = TestLogger(self.log_dir)
        self.fixture = LTBTestFixture()

    def _run_test(self, test_id, category, name, description, action, fn):
        """Execute a single test function, catching all exceptions."""
        t0 = time.time()
        try:
            result = fn()
            t1 = time.time()
            if result is None:
                result = {}
            self.logger.record(
                test_id=test_id, category=category, name=name,
                description=description, action=action,
                status=result.get("status", "PASS"),
                detail=result.get("detail", ""),
                files_generated=result.get("files_generated"),
                started=t0, finished=t1,
            )
        except Exception as e:
            t1 = time.time()
            self.logger.record(
                test_id=test_id, category=category, name=name,
                description=description, action=action,
                status="FAIL",
                detail=f"{type(e).__name__}: {e}\n{traceback.format_exc()[-300:]}",
                started=t0, finished=t1,
            )

    def run_all(self):
        self.fixture.setup()
        try:
            self._run_smoke_tests()
            self._run_upload_tests()
            self._run_generation_tests()
            self._run_evidence_split_tests()
            self._run_evidence_aggregate_tests()
            if self.include_slow:
                self._run_evidence_semantic_tests()
            else:
                self.logger._log("\n--- Category 6: Semantic Tests (SKIPPED — use --include-slow) ---")
        finally:
            self.fixture.teardown()

        return self.logger.finish()

    # ── Category 1: Smoke Tests ──────────────────────────────────────────────

    def _run_smoke_tests(self):
        self.logger._log("\n--- Category 1: Smoke Tests ---")
        c = self.fixture.client

        def sm001():
            r = c.get("/")
            assert r.status_code == 200, f"Expected 200, got {r.status_code}"
            html = r.data.decode()
            assert "LTB" in html or "Evidence" in html or "Brief" in html, "Page missing expected content"
            return {"detail": f"status={r.status_code} content_length={len(html)}"}

        self._run_test("SM-001", "smoke", "GET index page",
                       "GET / returns 200 with LTB content",
                       "GET /", sm001)

        def sm002():
            r = c.get("/evidence-stats")
            assert r.status_code == 200, f"Expected 200, got {r.status_code}"
            data = r.get_json()
            assert "N4" in data, "Response missing N4 key"
            return {"detail": f"keys={list(data.keys())}"}

        self._run_test("SM-002", "smoke", "GET evidence-stats",
                       "GET /evidence-stats returns JSON with N4 key",
                       "GET /evidence-stats", sm002)

        def sm003():
            stats_path = os.path.join(PROJECT_DIR, "evidence_stats.json")
            renamed = stats_path + ".bak_test"
            had_file = os.path.exists(stats_path)
            try:
                if had_file:
                    os.rename(stats_path, renamed)
                r = c.get("/evidence-stats")
                assert r.status_code == 200
                data = r.get_json()
                assert "N4" in data, "Fallback missing N4"
                assert data["N4"]["description"] == "Non-payment of rent"
                return {"detail": "fallback defaults returned correctly"}
            finally:
                if had_file and os.path.exists(renamed):
                    os.rename(renamed, stats_path)

        self._run_test("SM-003", "smoke", "GET evidence-stats fallback",
                       "With evidence_stats.json missing, defaults are returned",
                       "GET /evidence-stats (file removed)", sm003)

        def sm004():
            r = c.get("/thumbnail/nonexistent_abc123.jpg")
            assert r.status_code == 404, f"Expected 404, got {r.status_code}"
            return {"detail": f"status={r.status_code}"}

        self._run_test("SM-004", "smoke", "GET thumbnail 404",
                       "GET /thumbnail/nonexistent → 404",
                       "GET /thumbnail/nonexistent.jpg", sm004)

        def sm005():
            r = c.post("/generate", data="not json", content_type="text/plain")
            # Flask returns 415 (Unsupported Media Type) when content-type is not JSON
            # and request.get_json() is called, OR 400 if it parses as None
            assert r.status_code in (400, 415), f"Expected 400 or 415, got {r.status_code}"
            return {"detail": f"status={r.status_code}"}

        self._run_test("SM-005", "smoke", "POST generate no JSON body",
                       "POST /generate without JSON content-type → 400/415 error",
                       "POST /generate (text/plain body)", sm005)

    # ── Category 2: Upload Tests ─────────────────────────────────────────────

    def _run_upload_tests(self):
        self.logger._log("\n--- Category 2: Upload Tests ---")
        fx = self.fixture

        def up001():
            r = fx.upload_file(make_png(400, 300), "unit_hallway_photo.png")
            assert r.status_code == 200
            data = r.get_json()
            f = data["files"][0]
            assert f["is_image"] is True, f"Expected is_image=True, got {f['is_image']}"
            assert f["thumb"] is not None, "Expected thumbnail"
            return {"detail": f"id={f['id']} is_image={f['is_image']} thumb={f['thumb']}",
                    "files_generated": [{"name": f["id"], "size_bytes": os.path.getsize(
                        os.path.join(fx.tmp_dir, f["id"]))}]}

        self._run_test("UP-001", "upload", "Upload PNG — unit hallway photo",
                       "POST /upload with fake hallway photo (door, window, labels)",
                       "POST /upload multipart files=(PNG, 'unit_hallway_photo.png')", up001)

        def up002():
            r = fx.upload_file(make_jpeg(), "kitchen_water_damage.jpg")
            assert r.status_code == 200
            data = r.get_json()
            f = data["files"][0]
            assert f["is_image"] is True
            assert f["thumb"] is not None
            return {"detail": f"id={f['id']} is_image={f['is_image']}"}

        self._run_test("UP-002", "upload", "Upload JPEG — water damage photo",
                       "POST /upload with fake kitchen ceiling water-damage photo",
                       "POST /upload multipart files=(JPEG, 'kitchen_water_damage.jpg')", up002)

        def up003():
            r = fx.upload_file(make_webp(), "broken_window_latch.webp")
            assert r.status_code == 200
            data = r.get_json()
            f = data["files"][0]
            assert f["is_image"] is True
            assert f["thumb"] is not None
            return {"detail": f"id={f['id']} is_image={f['is_image']}"}

        self._run_test("UP-003", "upload", "Upload WEBP — broken latch photo",
                       "POST /upload with fake broken-window-latch photo",
                       "POST /upload multipart files=(WEBP, 'broken_window_latch.webp')", up003)

        def up004():
            r = fx.upload_file(make_png(2000, 2000), "high_res_unit_photo.png")
            assert r.status_code == 200
            data = r.get_json()
            f = data["files"][0]
            assert f["is_image"] is True
            thumb_path = os.path.join(fx.tmp_dir, f["thumb"])
            with PILImage.open(thumb_path) as img:
                assert max(img.size) <= 300, f"Thumb too large: {img.size}"
            return {"detail": f"thumb_size={PILImage.open(thumb_path).size}"}

        self._run_test("UP-004", "upload", "Upload large PNG (2000x2000 unit photo)",
                       "POST /upload with high-res unit photo, verify thumb ≤ 300px",
                       "POST /upload multipart files=(PNG 2000x2000)", up004)

        def up005():
            r = fx.upload_file(make_png(400, 300, mode="RGBA"), "unit_photo_transparent.png")
            assert r.status_code == 200
            data = r.get_json()
            f = data["files"][0]
            assert f["is_image"] is True
            return {"detail": f"id={f['id']} is_image={f['is_image']}"}

        self._run_test("UP-005", "upload", "Upload PNG RGBA — transparent overlay",
                       "POST /upload with RGBA unit photo (transparency channel)",
                       "POST /upload multipart files=(RGBA PNG)", up005)

        def up006():
            files = [
                (make_png(300, 250), "hallway_photo_1.png"),
                (make_jpeg(600, 400), "bathroom_damage.jpg"),
                (make_webp(200, 200), "bedroom_window.webp"),
            ]
            r = fx.upload_files_multi(files)
            assert r.status_code == 200
            data = r.get_json()
            assert len(data["files"]) == 3, f"Expected 3 files, got {len(data['files'])}"
            return {"detail": f"uploaded {len(data['files'])} files (hallway, bathroom, bedroom)"}

        self._run_test("UP-006", "upload", "Multi-file upload (3 room photos)",
                       "POST /upload with hallway PNG + bathroom JPEG + bedroom WEBP",
                       "POST /upload multipart files=(3 mixed-format room photos)", up006)

        def up007():
            r = fx.upload_file(make_pdf_text(), "rent_payment_record.pdf")
            assert r.status_code == 200
            data = r.get_json()
            f = data["files"][0]
            assert f["is_image"] is False, f"Expected is_image=False for PDF"
            return {"detail": f"id={f['id']} is_image={f['is_image']}"}

        self._run_test("UP-007", "upload", "Upload PDF — rent payment record",
                       "POST /upload with fake rent ledger PDF (table of monthly payments)",
                       "POST /upload multipart files=(PDF, 'rent_payment_record.pdf')", up007)

        def up008():
            r = fx.upload_file(make_pdf_with_image(), "maintenance_inspection_report.pdf")
            assert r.status_code == 200
            data = r.get_json()
            f = data["files"][0]
            assert f["is_image"] is False
            return {"detail": f"id={f['id']} is_image={f['is_image']}"}

        self._run_test("UP-008", "upload", "Upload PDF — inspection report with photo",
                       "POST /upload with fake maintenance inspection report PDF (text + embedded photo)",
                       "POST /upload multipart files=(PDF, 'maintenance_inspection_report.pdf')", up008)

        def up009():
            buf = make_docx()
            if buf is None:
                return {"status": "SKIP", "detail": "python-docx not installed"}
            r = fx.upload_file(buf, "demand_letter_rent_arrears.docx")
            assert r.status_code == 200
            data = r.get_json()
            f = data["files"][0]
            assert f["is_image"] is False
            return {"detail": f"id={f['id']} is_image={f['is_image']}"}

        self._run_test("UP-009", "upload", "Upload DOCX — demand letter",
                       "POST /upload with fake rent-arrears demand letter DOCX",
                       "POST /upload multipart files=(DOCX, 'demand_letter_rent_arrears.docx')", up009)

        def up010():
            r = fx.upload_file(make_md(), "tenant_communication_log.md")
            assert r.status_code == 200
            data = r.get_json()
            f = data["files"][0]
            assert f["is_image"] is False
            return {"detail": f"id={f['id']} is_image={f['is_image']}"}

        self._run_test("UP-010", "upload", "Upload Markdown — communication log",
                       "POST /upload with fake tenant communication log (.md)",
                       "POST /upload multipart files=(MD, 'tenant_communication_log.md')", up010)

        def up011():
            r = fx.client.post("/upload", data={}, content_type="multipart/form-data")
            assert r.status_code == 400, f"Expected 400, got {r.status_code}"
            return {"detail": f"status={r.status_code}"}

        self._run_test("UP-011", "upload", "No files field",
                       "POST /upload with empty multipart body → 400",
                       "POST /upload (empty)", up011)

        def up012():
            bad_buf = io.BytesIO(b"\xff\xd8\xff\x00NOTAJPEG\x00\x00")
            bad_buf.name = "corrupt.jpg"
            r = fx.upload_file(bad_buf, "corrupt.jpg")
            assert r.status_code == 200, f"Expected 200, got {r.status_code}"
            data = r.get_json()
            f = data["files"][0]
            # Corrupted image can't be opened by PIL so is_image should be False
            assert f["is_image"] is False, f"Expected is_image=False for corrupt JPEG"
            return {"detail": f"id={f['id']} is_image={f['is_image']}"}

        self._run_test("UP-012", "upload", "Corrupted JPEG upload",
                       "POST /upload with invalid bytes as .jpg → 200, is_image=False",
                       "POST /upload multipart files=(corrupt JPEG)", up012)

        def up013():
            r = fx.upload_file(make_gif(), "front_door_damage.gif")
            assert r.status_code == 200
            data = r.get_json()
            f = data["files"][0]
            assert f["is_image"] is True, f"Expected is_image=True for GIF"
            assert f["thumb"] is not None
            return {"detail": f"id={f['id']} is_image={f['is_image']}"}

        self._run_test("UP-013", "upload", "Upload GIF — door damage photo",
                       "POST /upload with fake scratched front door photo (GIF)",
                       "POST /upload multipart files=(GIF, 'front_door_damage.gif')", up013)

        def up014():
            r = fx.upload_file(make_bmp(), "bathroom_tile_cracked.bmp")
            assert r.status_code == 200
            data = r.get_json()
            f = data["files"][0]
            assert f["is_image"] is True, f"Expected is_image=True for BMP"
            assert f["thumb"] is not None
            return {"detail": f"id={f['id']} is_image={f['is_image']}"}

        self._run_test("UP-014", "upload", "Upload BMP — cracked tile photo",
                       "POST /upload with fake cracked bathroom tile photo (BMP)",
                       "POST /upload multipart files=(BMP, 'bathroom_tile_cracked.bmp')", up014)

        def up015():
            r = fx.upload_file(make_tiff(), "bedroom_mold_growth.tiff")
            assert r.status_code == 200
            data = r.get_json()
            f = data["files"][0]
            assert f["is_image"] is True, f"Expected is_image=True for TIFF"
            assert f["thumb"] is not None
            return {"detail": f"id={f['id']} is_image={f['is_image']}"}

        self._run_test("UP-015", "upload", "Upload TIFF — mold growth photo",
                       "POST /upload with fake bedroom mold photo (TIFF)",
                       "POST /upload multipart files=(TIFF, 'bedroom_mold_growth.tiff')", up015)

        def up016():
            buf = make_heic()
            if buf is None:
                return {"status": "SKIP", "detail": "pillow-heif not installed"}
            r = fx.upload_file(buf, "carpet_stain.heic")
            assert r.status_code == 200
            data = r.get_json()
            f = data["files"][0]
            # HEIC gets converted to JPEG by the app, so it should be is_image=True
            assert f["is_image"] is True, f"Expected is_image=True for HEIC (converted)"
            return {"detail": f"id={f['id']} is_image={f['is_image']} (HEIC→JPEG converted)"}

        self._run_test("UP-016", "upload", "Upload HEIC — carpet stain photo",
                       "POST /upload with fake carpet stain photo (HEIC), app converts to JPEG",
                       "POST /upload multipart files=(HEIC, 'carpet_stain.heic')", up016)

    # ── Category 3: PDF Generation Tests ─────────────────────────────────────

    def _run_generation_tests(self):
        self.logger._log("\n--- Category 3: PDF Generation Tests ---")
        fx = self.fixture

        def _basic_case_info():
            return {
                "file_number": "TSL-99999-26",
                "applicant_name": "ABC Property Management Inc.",
                "respondent_name": "Jane Doe",
                "respondent_address": "123 Test Street, Unit 4B, Toronto ON M5V 2T6",
            }

        def _upload_and_get_id(file_buf, filename):
            """Upload a file and return its server-side ID."""
            r = fx.upload_file(file_buf, filename)
            data = r.get_json()
            return data["files"][0]["id"]

        def gen001():
            r = fx.generate_brief(_basic_case_info(), [{"title": "Empty Tab", "images": []}])
            assert r.status_code == 200, f"Expected 200, got {r.status_code}"
            pdf_data = r.data
            assert pdf_data[:5] == b"%PDF-", "Response is not a PDF"
            assert len(pdf_data) > 1024, f"PDF too small: {len(pdf_data)} bytes"
            return {"detail": f"pdf_size={len(pdf_data)} bytes",
                    "files_generated": [{"name": "brief.pdf", "size_bytes": len(pdf_data)}]}

        self._run_test("GEN-001", "generation", "Minimal brief (empty tab)",
                       "Case info + 1 empty tab → valid PDF > 1KB",
                       "POST /generate (minimal)", gen001)

        def gen002():
            img_id = _upload_and_get_id(make_jpeg(), "kitchen_water_damage.jpg")
            r = fx.generate_brief(_basic_case_info(), [{"title": "Unit Condition Photos", "images": [img_id]}])
            assert r.status_code == 200
            pdf_data = r.data
            assert pdf_data[:5] == b"%PDF-"
            assert "attachment" in r.headers.get("Content-Disposition", "").lower()
            return {"detail": f"pdf_size={len(pdf_data)} has_disposition=True",
                    "files_generated": [{"name": "brief.pdf", "size_bytes": len(pdf_data)}]}

        self._run_test("GEN-002", "generation", "Single photo N4 brief",
                       "1 tab with kitchen damage photo → valid PDF with Content-Disposition",
                       "POST /generate (1 tab, 1 damage photo)", gen002)

        def gen003():
            photo_names = ["kitchen_ceiling.jpg", "bathroom_mold.jpg", "bedroom_window.jpg"]
            ids = [_upload_and_get_id(make_jpeg(400 + i*100, 300 + i*50), n) for i, n in enumerate(photo_names)]
            r = fx.generate_brief(_basic_case_info(), [{"title": "Unit Damage Photos", "images": ids}])
            assert r.status_code == 200
            pdf_data = r.data
            assert pdf_data[:5] == b"%PDF-"
            return {"detail": f"pdf_size={len(pdf_data)} (3 damage photos)",
                    "files_generated": [{"name": "brief.pdf", "size_bytes": len(pdf_data)}]}

        self._run_test("GEN-003", "generation", "Multi-photo tab (3 damage photos)",
                       "1 tab, 3 damage photos (kitchen, bathroom, bedroom) → valid PDF",
                       "POST /generate (1 tab, 3 damage photos)", gen003)

        def gen004():
            tab_names = ["N4 Notice", "Financial Records", "Communication"]
            tabs = []
            for i, title in enumerate(tab_names):
                img_id = _upload_and_get_id(make_png(300, 250), f"evidence_tab{i+1}.png")
                tabs.append({"title": title, "images": [img_id]})
            r = fx.generate_brief(_basic_case_info(), tabs)
            assert r.status_code == 200
            assert r.data[:5] == b"%PDF-"
            return {"detail": f"pdf_size={len(r.data)} tabs={tab_names}"}

        self._run_test("GEN-004", "generation", "Three evidence tabs N4",
                       "3 tabs (N4 Notice, Financial, Communication) → valid PDF",
                       "POST /generate (3 evidence tabs)", gen004)

        def gen005():
            tab_configs = [
                ("N4 Notice", ["n4_notice_page1.png", "n4_notice_page2.png"]),
                ("Rent Ledger", ["rent_ledger_jan.png", "rent_ledger_feb.png"]),
                ("Lease Agreement", ["lease_page1.png", "lease_page2.png"]),
                ("Communication", ["email_screenshot1.png", "text_messages.png"]),
                ("Unit Photos", ["kitchen_damage.png", "hallway_photo.png"]),
            ]
            tabs = []
            for title, filenames in tab_configs:
                ids = [_upload_and_get_id(make_png(300 + j*50, 400 + j*30), fn) for j, fn in enumerate(filenames)]
                tabs.append({"title": title, "images": ids})
            r = fx.generate_brief(_basic_case_info(), tabs)
            assert r.status_code == 200
            pdf_data = r.data
            assert pdf_data[:5] == b"%PDF-"
            assert len(pdf_data) > 10000, f"PDF too small for 5 tabs: {len(pdf_data)}"
            return {"detail": f"pdf_size={len(pdf_data)} (5 tabs, 10 evidence files)"}

        self._run_test("GEN-005", "generation", "Five tabs stress (full evidence brief)",
                       "5 tabs (Notice, Ledger, Lease, Comms, Photos), 2 files each → valid PDF > 10KB",
                       "POST /generate (5 tabs, 10 files)", gen005)

        def gen006():
            img_id = _upload_and_get_id(make_jpeg(), "maintenance_issue_photo.jpg")
            case = _basic_case_info()
            case["file_number"] = "SOL-55555-26"
            tabs = [
                {"title": "Maintenance Issues", "images": [img_id]},
                {"title": "Landlord Communication", "images": []},
            ]
            r = fx.generate_brief(case, tabs)
            assert r.status_code == 200
            assert r.data[:5] == b"%PDF-"
            return {"detail": f"pdf_size={len(r.data)} file_number=SOL-55555-26"}

        self._run_test("GEN-006", "generation", "N5 application brief",
                       "N5 maintenance case (SOL-55555-26) with 2 tabs → valid PDF",
                       "POST /generate (N5, 2 tabs)", gen006)

        def gen007():
            buf = make_docx()
            if buf is None:
                return {"status": "SKIP", "detail": "python-docx not installed"}
            doc_id = _upload_and_get_id(buf, "demand_letter_rent_arrears.docx")
            r = fx.generate_brief(_basic_case_info(), [{"title": "Demand Letter", "images": [doc_id]}])
            assert r.status_code == 200
            assert r.data[:5] == b"%PDF-"
            return {"detail": f"pdf_size={len(r.data)}"}

        self._run_test("GEN-007", "generation", "DOCX demand letter in tab",
                       "Upload demand letter DOCX, embed in brief → valid PDF",
                       "POST /generate (demand letter DOCX in tab)", gen007)

        def gen008():
            pdf_id = _upload_and_get_id(make_pdf_text(), "rent_payment_record.pdf")
            r = fx.generate_brief(_basic_case_info(), [{"title": "Payment Records", "images": [pdf_id]}])
            assert r.status_code == 200
            assert r.data[:5] == b"%PDF-"
            return {"detail": f"pdf_size={len(r.data)}"}

        self._run_test("GEN-008", "generation", "Rent payment PDF in tab",
                       "Upload rent payment record PDF, embed in brief → valid PDF",
                       "POST /generate (rent record PDF in tab)", gen008)

        def gen009():
            r = fx.generate_brief(
                _basic_case_info(),
                [{"title": "Bad Refs", "images": ["nonexistent_abc.png", "fake_xyz.pdf"]}]
            )
            assert r.status_code == 200, f"Expected 200, got {r.status_code}"
            assert r.data[:5] == b"%PDF-"
            return {"detail": f"pdf_size={len(r.data)} (gracefully handled bad IDs)"}

        self._run_test("GEN-009", "generation", "Bad file IDs in tab",
                       "Nonexistent file IDs → 200, valid PDF (graceful)",
                       "POST /generate (bad file IDs)", gen009)

        def gen010():
            img_id = _upload_and_get_id(make_gif(), "front_door_damage.gif")
            r = fx.generate_brief(_basic_case_info(), [{"title": "Door Damage Photos", "images": [img_id]}])
            assert r.status_code == 200
            assert r.data[:5] == b"%PDF-"
            return {"detail": f"pdf_size={len(r.data)}"}

        self._run_test("GEN-010", "generation", "GIF image in brief",
                       "Upload GIF door damage photo, embed in brief → valid PDF",
                       "POST /generate (GIF in tab)", gen010)

        def gen011():
            img_id = _upload_and_get_id(make_bmp(), "bathroom_tile_cracked.bmp")
            r = fx.generate_brief(_basic_case_info(), [{"title": "Tile Damage", "images": [img_id]}])
            assert r.status_code == 200
            assert r.data[:5] == b"%PDF-"
            return {"detail": f"pdf_size={len(r.data)}"}

        self._run_test("GEN-011", "generation", "BMP image in brief",
                       "Upload BMP cracked tile photo, embed in brief → valid PDF",
                       "POST /generate (BMP in tab)", gen011)

        def gen012():
            img_id = _upload_and_get_id(make_tiff(), "bedroom_mold_growth.tiff")
            r = fx.generate_brief(_basic_case_info(), [{"title": "Mold Evidence", "images": [img_id]}])
            assert r.status_code == 200
            assert r.data[:5] == b"%PDF-"
            return {"detail": f"pdf_size={len(r.data)}"}

        self._run_test("GEN-012", "generation", "TIFF image in brief",
                       "Upload TIFF mold photo, embed in brief → valid PDF",
                       "POST /generate (TIFF in tab)", gen012)

        def gen013():
            md_id = _upload_and_get_id(make_md(), "tenant_communication_log.md")
            r = fx.generate_brief(_basic_case_info(), [{"title": "Communication Log", "images": [md_id]}])
            assert r.status_code == 200
            assert r.data[:5] == b"%PDF-"
            return {"detail": f"pdf_size={len(r.data)}"}

        self._run_test("GEN-013", "generation", "Markdown file in brief",
                       "Upload Markdown communication log, embed in brief → valid PDF",
                       "POST /generate (MD in tab)", gen013)

        def gen014():
            buf = make_heic()
            if buf is None:
                return {"status": "SKIP", "detail": "pillow-heif not installed"}
            heic_id = _upload_and_get_id(buf, "carpet_stain.heic")
            r = fx.generate_brief(_basic_case_info(), [{"title": "Carpet Stain", "images": [heic_id]}])
            assert r.status_code == 200
            assert r.data[:5] == b"%PDF-"
            return {"detail": f"pdf_size={len(r.data)}"}

        self._run_test("GEN-014", "generation", "HEIC image in brief",
                       "Upload HEIC carpet stain photo (converted to JPEG), embed in brief → valid PDF",
                       "POST /generate (HEIC in tab)", gen014)

        def gen015():
            # Upload one of every supported format
            jpeg_id = _upload_and_get_id(make_jpeg(500, 400), "kitchen_damage.jpg")
            png_id = _upload_and_get_id(make_png(300, 250), "hallway_evidence.png")
            webp_id = _upload_and_get_id(make_webp(200, 200), "window_latch.webp")
            gif_id = _upload_and_get_id(make_gif(300, 250), "door_scratches.gif")
            bmp_id = _upload_and_get_id(make_bmp(300, 250), "tile_crack.bmp")
            tiff_id = _upload_and_get_id(make_tiff(300, 250), "wall_mold.tiff")
            pdf_id = _upload_and_get_id(make_pdf_text(), "rent_record.pdf")
            md_id = _upload_and_get_id(make_md(), "comm_log.md")

            tab_list = [
                {"title": "Unit Photos (JPEG + PNG + WEBP)", "images": [jpeg_id, png_id, webp_id]},
                {"title": "Additional Photos (GIF + BMP + TIFF)", "images": [gif_id, bmp_id, tiff_id]},
                {"title": "Rent Payment Records (PDF)", "images": [pdf_id]},
                {"title": "Communication Log (Markdown)", "images": [md_id]},
            ]

            docx_buf = make_docx()
            if docx_buf is not None:
                docx_id = _upload_and_get_id(docx_buf, "demand_letter.docx")
                tab_list.append({"title": "Demand Letter (DOCX)", "images": [docx_id]})

            heic_buf = make_heic()
            if heic_buf is not None:
                heic_id = _upload_and_get_id(heic_buf, "carpet_photo.heic")
                tab_list[1]["images"].append(heic_id)  # Add to photos tab

            r = fx.generate_brief(_basic_case_info(), tab_list)
            assert r.status_code == 200
            pdf_data = r.data
            assert pdf_data[:5] == b"%PDF-"
            fmt_count = 6 + (1 if docx_buf else 0) + (1 if heic_buf else 0)
            return {"detail": f"pdf_size={len(pdf_data)} formats_used={fmt_count} tabs={len(tab_list)}"}

        self._run_test("GEN-015", "generation", "All-formats mixed brief",
                       "Every supported format (JPEG+PNG+WEBP+GIF+BMP+TIFF+PDF+MD+DOCX+HEIC) in one brief",
                       "POST /generate (all formats combined)", gen015)

    # ── Category 4: Evidence Analyzer — Case Splitting ───────────────────────

    def _run_evidence_split_tests(self):
        self.logger._log("\n--- Category 4: Evidence Analyzer — Case Splitting ---")
        from evidence_analyzer import _split_into_cases, _chunk_text

        def ea001():
            result = _split_into_cases("This is plain text with no case numbers at all.")
            assert len(result) == 1, f"Expected 1 block, got {len(result)}"
            assert result[0][0] == "FULL_DOC"
            return {"detail": f"blocks={len(result)} id={result[0][0]}"}

        self._run_test("EA-001", "evidence-split", "No case numbers",
                       "_split_into_cases with plain text → [('FULL_DOC', ...)]",
                       "_split_into_cases(plain text)", ea001)

        def ea002():
            text = "File Number: TSL-12345-22\nThe tenant failed to pay rent for three consecutive months. " \
                   "The landlord served an N4 notice of termination on the tenant."
            result = _split_into_cases(text)
            assert len(result) == 1, f"Expected 1 case, got {len(result)}"
            assert result[0][0] == "TSL-12345-22"
            return {"detail": f"cases={len(result)} id={result[0][0]}"}

        self._run_test("EA-002", "evidence-split", "Single case number",
                       "_split_into_cases with one TSL file number → 1 case",
                       "_split_into_cases(single case)", ea002)

        def ea003():
            text = ("File Number: TSL-12345-22\n"
                    "The tenant owed rent. The landlord served an N4 notice. "
                    "The Board found the tenant breached the agreement.\n\n"
                    "File Number: SOL-98765-23\n"
                    "A different tenancy dispute. The tenant failed to maintain the unit. "
                    "The landlord provided photos as evidence of damage.")
            result = _split_into_cases(text)
            assert len(result) == 2, f"Expected 2 cases, got {len(result)}"
            ids = [r[0] for r in result]
            assert "TSL-12345-22" in ids
            assert "SOL-98765-23" in ids
            return {"detail": f"cases={len(result)} ids={ids}"}

        self._run_test("EA-003", "evidence-split", "Two case numbers",
                       "_split_into_cases with two file numbers → 2 cases",
                       "_split_into_cases(two cases)", ea003)

        def ea004():
            text = "File Number: TSL-11111-22\nToo short.\nFile Number: TSL-22222-22\n" + ("x " * 40)
            result = _split_into_cases(text)
            # First block "Too short." is < 50 chars; second block is long enough
            short_ids = [r[0] for r in result if len(r[1].strip()) <= 50]
            assert len(short_ids) == 0, f"Small blocks should be dropped: {short_ids}"
            return {"detail": f"cases={len(result)} (small blocks dropped)"}

        self._run_test("EA-004", "evidence-split", "Block too small (< 50 chars)",
                       "_split_into_cases drops blocks shorter than 50 chars",
                       "_split_into_cases(small block)", ea004)

        def ea005():
            text = ("File Number: TSL-12345-22\n"
                    "First occurrence with enough text to be meaningful for analysis purposes.\n\n"
                    "File Number: TSL-12345-22\n"
                    "Second occurrence of the same case number should be ignored entirely.")
            result = _split_into_cases(text)
            ids = [r[0] for r in result]
            assert ids.count("TSL-12345-22") == 1, f"Duplicate IDs: {ids}"
            return {"detail": f"cases={len(result)} (no duplicates)"}

        self._run_test("EA-005", "evidence-split", "Duplicate case IDs",
                       "_split_into_cases with same ID twice → only first kept",
                       "_split_into_cases(duplicate IDs)", ea005)

        def ea006():
            text = ("File Number: LTB-L-30001-25\n"
                    "An AI-generated LTB decision. The landlord served the tenant with an N4 notice "
                    "of termination for non-payment of rent arrears totaling several months.")
            result = _split_into_cases(text)
            assert len(result) == 1, f"Expected 1 case, got {len(result)}"
            assert result[0][0] == "LTB-L-30001-25"
            return {"detail": f"id={result[0][0]} (AI-generated format matched)"}

        self._run_test("EA-006", "evidence-split", "AI-generated case ID",
                       "_split_into_cases matches LTB-L-30001-25 format",
                       "_split_into_cases(AI ID)", ea006)

        def ea007():
            text = ("First sentence about evidence. Second sentence about the hearing. "
                    "Third sentence about the decision. Fourth sentence about remedies. "
                    "Fifth sentence about the order. Sixth sentence is the conclusion.")
            chunks = _chunk_text(text, max_chunk_len=200)
            assert len(chunks) >= 1, "Expected at least 1 chunk"
            for ch in chunks:
                assert len(ch) <= 300, f"Chunk too long: {len(ch)} chars"
            return {"detail": f"chunks={len(chunks)} max_len={max(len(c) for c in chunks)}"}

        self._run_test("EA-007", "evidence-split", "Normal chunking",
                       "_chunk_text splits paragraph into chunks ≤ ~200 chars",
                       "_chunk_text(paragraph)", ea007)

        def ea008():
            chunks = _chunk_text("")
            assert chunks == [], f"Expected [], got {chunks}"
            return {"detail": "empty string → []"}

        self._run_test("EA-008", "evidence-split", "Empty string chunking",
                       "_chunk_text('') → []",
                       "_chunk_text('')", ea008)

    # ── Category 5: Evidence Analyzer — Aggregation ──────────────────────────

    def _run_evidence_aggregate_tests(self):
        self.logger._log("\n--- Category 5: Evidence Analyzer — Aggregation ---")
        from evidence_analyzer import aggregate_stats

        def ea009():
            result = aggregate_stats([])
            assert result["N4"]["total_cases_analyzed"] == 0
            assert result["N4"]["evidence_types"] == []
            return {"detail": "empty input → total=0, empty types"}

        self._run_test("EA-009", "evidence-agg", "Aggregate empty list",
                       "aggregate_stats([]) → total=0, empty evidence_types",
                       "aggregate_stats([])", ea009)

        def ea010():
            cases = [
                {"case_id": "A", "evidence_types": ["N4 Notice of Termination", "Lease agreement"]},
                {"case_id": "B", "evidence_types": ["N4 Notice of Termination"]},
                {"case_id": "C", "evidence_types": ["Lease agreement"]},
            ]
            result = aggregate_stats(cases)
            stats = result["N4"]
            assert stats["total_cases_analyzed"] == 3
            pcts = {e["category"]: e["percentage"] for e in stats["evidence_types"]}
            assert pcts["N4 Notice of Termination"] == 67, f"Expected 67%, got {pcts.get('N4 Notice of Termination')}"
            assert pcts["Lease agreement"] == 67, f"Expected 67%, got {pcts.get('Lease agreement')}"
            return {"detail": f"total=3 N4_pct={pcts.get('N4 Notice of Termination')} lease_pct={pcts.get('Lease agreement')}"}

        self._run_test("EA-010", "evidence-agg", "Aggregate normal (3 cases)",
                       "3 cases → correct percentages (67%)",
                       "aggregate_stats(3 cases)", ea010)

        def ea011():
            cases = [{"case_id": str(i), "evidence_types": ["Witness testimony"]} for i in range(5)]
            result = aggregate_stats(cases)
            stats = result["N4"]
            assert stats["total_cases_analyzed"] == 5
            pcts = {e["category"]: e["percentage"] for e in stats["evidence_types"]}
            assert pcts["Witness testimony"] == 100
            return {"detail": "5 identical cases → 100%"}

        self._run_test("EA-011", "evidence-agg", "Aggregate uniform (5 identical)",
                       "5 identical cases → 100% for that type",
                       "aggregate_stats(5 identical)", ea011)

    # ── Category 6: Evidence Analyzer — Semantic (slow) ──────────────────────

    def _run_evidence_semantic_tests(self):
        self.logger._log("\n--- Category 6: Semantic Tests (slow) ---")

        def ea012():
            from evidence_analyzer import analyze_cases
            text = (
                "File Number: TSL-99999-26\n"
                "The landlord served an N4 notice of termination on the tenant for non-payment "
                "of rent. The tenant owed rent arrears totaling $5,000. The landlord submitted "
                "bank statements and a rent ledger as evidence. The lease agreement was entered "
                "into evidence at the hearing. The Board ordered termination of the tenancy."
            )
            results = analyze_cases(text, similarity_threshold=0.45, verbose=False)
            assert len(results) >= 1, "Expected at least 1 case"
            types = results[0]["evidence_types"]
            assert "N4 Notice of Termination" in types, f"Expected N4 Notice detected, got: {types}"
            return {"detail": f"detected_types={types}"}

        self._run_test("EA-012", "evidence-semantic", "Similarity hit (N4 text)",
                       "N4 decision text → 'N4 Notice' category detected",
                       "analyze_cases(N4 decision text)", ea012)

        def ea013():
            from evidence_analyzer import analyze_cases
            text = (
                "Today we will make a delicious pasta carbonara. Start by boiling water in a large pot. "
                "Cook the spaghetti until al dente. In a separate pan, fry the pancetta until crispy. "
                "Mix egg yolks with grated pecorino cheese. Combine everything and serve immediately. "
                "This recipe serves four people and takes about thirty minutes to prepare."
            )
            results = analyze_cases(text, similarity_threshold=0.50, verbose=False)
            if results:
                types = results[0]["evidence_types"]
                assert len(types) == 0, f"Expected no categories for cooking recipe, got: {types}"
            return {"detail": f"categories_found={results[0]['evidence_types'] if results else 'none'}"}

        self._run_test("EA-013", "evidence-semantic", "Similarity miss (cooking recipe)",
                       "Unrelated text → no evidence categories detected",
                       "analyze_cases(cooking recipe)", ea013)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="LTB Evidence Brief Generator — Test Suite")
    parser.add_argument("--include-slow", action="store_true",
                        help="Include slow semantic model tests (~2 min)")
    args = parser.parse_args()

    runner = TestRunner(include_slow=args.include_slow)
    failures = runner.run_all()
    sys.exit(1 if failures > 0 else 0)


if __name__ == "__main__":
    main()
