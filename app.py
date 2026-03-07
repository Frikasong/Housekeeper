# -*- coding: utf-8 -*-
import os
import json
import uuid
import io
import time
import secrets
import subprocess
from xml.sax.saxutils import escape as xml_escape
from flask import Flask, request, jsonify, send_file, render_template
from werkzeug.utils import secure_filename
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak,
    Image as RLImage, Table, TableStyle, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.pdfgen import canvas
from reportlab.platypus.flowables import Flowable

# Optional text-extraction dependencies (app works without them)
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import fitz as pymupdf
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import pillow_heif
    pillow_heif.register_heif_opener()
    HAS_PILLOW_HEIF = True
except ImportError:
    HAS_PILLOW_HEIF = False

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', secrets.token_hex(32))
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB

IMAGE_EXTS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp'}
PLAIN_TEXT_EXTS = {
    'txt', 'md', 'csv', 'log', 'py', 'js', 'ts', 'html', 'htm',
    'css', 'json', 'xml', 'yaml', 'yml', 'ini', 'cfg', 'conf',
    'sql', 'sh', 'bat', 'tex', 'rst', 'rtf',
}
OFFICE_EXTS = {'docx', 'doc', 'xlsx', 'xls', 'pptx', 'ppt',
               'odt', 'ods', 'odp'}
HEIC_EXTS = {'heic', 'heif'}

def is_image_file(path):
    ext = path.rsplit('.', 1)[-1].lower() if '.' in path else ''
    return ext in IMAGE_EXTS


def extract_text_from_file(path):
    """Return (text, success). Reads the full extractable text with no truncation."""
    ext = path.rsplit('.', 1)[-1].lower() if '.' in path else ''
    try:
        if ext in PLAIN_TEXT_EXTS:
            with open(path, 'r', encoding='utf-8', errors='replace') as f:
                raw = f.read()
            return raw, True

        elif ext == 'pdf':
            if not HAS_PDFPLUMBER:
                return None, False
            parts = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    parts.append(page.extract_text() or '')
            return '\n'.join(parts), True

        elif ext == 'docx':
            if not HAS_DOCX:
                return None, False
            doc = DocxDocument(path)
            raw = '\n'.join(p.text for p in doc.paragraphs if p.text)
            return raw, True

        elif ext in ('xlsx', 'xls'):
            if not HAS_OPENPYXL:
                return None, False
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append("── Sheet: %s ──" % sheet_name)
                for row in ws.iter_rows(values_only=True):
                    cells_list = [str(c) if c is not None else '' for c in row]
                    parts.append('\t'.join(cells_list))
            wb.close()
            return '\n'.join(parts), True

        else:
            return None, False

    except Exception:
        return None, False


# ─── DOCUMENT-TO-PDF CONVERSION ───────────────────────────────────────────────

def _find_libreoffice():
    """Locate LibreOffice / soffice executable."""
    for cmd in ['soffice', 'libreoffice',
                '/Applications/LibreOffice.app/Contents/MacOS/soffice']:
        try:
            subprocess.run([cmd, '--version'], capture_output=True, timeout=10)
            return cmd
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    return None

_LO_CMD_CACHE = [None, False]   # [command, searched]

def _get_lo_cmd():
    if not _LO_CMD_CACHE[1]:
        _LO_CMD_CACHE[0] = _find_libreoffice()
        _LO_CMD_CACHE[1] = True
    return _LO_CMD_CACHE[0]


def _convert_heic_to_jpeg(src_path, dest_path):
    """Convert a HEIC/HEIF file to JPEG.
    Tries pillow-heif first; falls back to macOS sips."""
    if HAS_PILLOW_HEIF:
        try:
            with Image.open(src_path) as img:
                img.convert('RGB').save(dest_path, 'JPEG', quality=95)
            return os.path.exists(dest_path)
        except Exception:
            pass
    try:
        r = subprocess.run(
            ['sips', '-s', 'format', 'jpeg', src_path, '--out', dest_path],
            capture_output=True, timeout=60,
        )
        return r.returncode == 0 and os.path.exists(dest_path)
    except Exception:
        return False


def _convert_office_to_pdf(file_path):
    """Convert an office document to PDF via LibreOffice headless."""
    lo_cmd = _get_lo_cmd()
    if not lo_cmd:
        return None
    try:
        outdir = app.config['UPLOAD_FOLDER']
        subprocess.run(
            [lo_cmd, '--headless', '--convert-to', 'pdf',
             '--outdir', outdir, file_path],
            capture_output=True, timeout=120,
        )
        base = os.path.splitext(os.path.basename(file_path))[0]
        out_path = os.path.join(outdir, base + '.pdf')
        if os.path.exists(out_path):
            new_path = os.path.join(outdir, 'conv_%s.pdf' % uuid.uuid4().hex)
            os.rename(out_path, new_path)
            return new_path
    except Exception:
        pass
    return None


def _convert_text_to_pdf(file_path):
    """Render a plain-text file as a PDF using ReportLab. Returns path or None."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            text = f.read()
    except Exception:
        return None
    if not text.strip():
        return None
    out_path = os.path.join(
        app.config['UPLOAD_FOLDER'],
        'textconv_%s.pdf' % uuid.uuid4().hex,
    )
    try:
        doc = SimpleDocTemplate(
            out_path, pagesize=letter,
            leftMargin=0.6 * inch, rightMargin=0.6 * inch,
            topMargin=0.6 * inch, bottomMargin=0.6 * inch,
        )
        styles = getSampleStyleSheet()
        code_style = ParagraphStyle(
            'TextConv', parent=styles['Normal'],
            fontName='Courier', fontSize=8, leading=10,
            spaceBefore=0, spaceAfter=0,
        )
        story = []
        for line in text.split('\n'):
            safe = xml_escape(line) if line.strip() else '&nbsp;'
            story.append(Paragraph(safe, code_style))
        doc.build(story)
        return out_path
    except Exception:
        try:
            os.unlink(out_path)
        except OSError:
            pass
        return None


def _convert_docx_to_pdf_reportlab(file_path):
    """Convert .docx to PDF using python-docx, preserving text and embedded images.
    Walks the document body XML in order so images appear in context with surrounding text."""
    if not HAS_DOCX:
        return None
    try:
        doc = DocxDocument(file_path)
    except Exception:
        return None

    # XML namespaces used in docx/drawingml
    _A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    _R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    out_path = os.path.join(
        app.config['UPLOAD_FOLDER'],
        'docxconv_%s.pdf' % uuid.uuid4().hex,
    )

    rl_styles = getSampleStyleSheet()
    normal_s = ParagraphStyle('DxBody', parent=rl_styles['Normal'],
                              fontName='Helvetica', fontSize=10, leading=14, spaceAfter=3)
    h1_s = ParagraphStyle('DxH1', parent=rl_styles['Normal'],
                           fontName='Helvetica-Bold', fontSize=14, leading=18,
                           spaceBefore=10, spaceAfter=4)
    h2_s = ParagraphStyle('DxH2', parent=rl_styles['Normal'],
                           fontName='Helvetica-Bold', fontSize=12, leading=16,
                           spaceBefore=8, spaceAfter=3)

    usable_w = letter[0] - 1.5 * inch

    def _image_flowable(blob):
        try:
            buf = io.BytesIO(blob)
            with Image.open(buf) as pil:
                orig_w, orig_h = pil.size
                mode = pil.mode
            # ReportLab struggles with RGBA/palette modes — normalise to RGB JPEG
            if mode in ('RGBA', 'P', 'LA'):
                buf.seek(0)
                out_buf = io.BytesIO()
                with Image.open(buf) as pil:
                    pil.convert('RGB').save(out_buf, 'JPEG', quality=90)
                out_buf.seek(0)
                buf = out_buf
            else:
                buf.seek(0)
            scale = min(usable_w / orig_w, 4 * inch / orig_h, 1.0)
            rl = RLImage(buf, width=orig_w * scale, height=orig_h * scale)
            rl.hAlign = 'CENTER'
            return rl
        except Exception:
            return None

    def _para_flowables(p_elem):
        """Return flowables for a single <w:p> element (images + text)."""
        items = []
        for blip in p_elem.iter('{%s}blip' % _A):
            rId = blip.get('{%s}embed' % _R)
            if rId:
                try:
                    blob = doc.part.related_parts[rId].blob
                    fl = _image_flowable(blob)
                    if fl:
                        items.append(fl)
                        items.append(Spacer(1, 0.1 * inch))
                except Exception:
                    pass
        texts = [t.text or '' for t in p_elem.iter('{%s}t' % _W)]
        text = ''.join(texts).strip()
        if text:
            ps_elem = p_elem.find('.//{%s}pStyle' % _W)
            sv = ps_elem.get('{%s}val' % _W, '') if ps_elem is not None else ''
            if 'Heading1' in sv or 'Title' in sv:
                ps = h1_s
            elif 'Heading' in sv:
                ps = h2_s
            else:
                ps = normal_s
            items.append(Paragraph(xml_escape(text), ps))
        return items

    try:
        story = []
        for child in doc.element.body:
            local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if local == 'p':
                story.extend(_para_flowables(child))
            elif local == 'tbl':
                for tr in child.iter('{%s}tr' % _W):
                    for tc in tr.findall('{%s}tc' % _W):
                        for p in tc.findall('{%s}p' % _W):
                            story.extend(_para_flowables(p))
        if not story:
            return None
        rl_doc = SimpleDocTemplate(
            out_path, pagesize=letter,
            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
            topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        )
        rl_doc.build(story)
        return out_path
    except Exception:
        try:
            os.unlink(out_path)
        except OSError:
            pass
        return None


def convert_document_to_images(file_path):
    """Convert ANY non-image file to rendered page images using PyMuPDF.

    Handles PDFs directly, office docs via LibreOffice, and text files
    via ReportLab. Returns a list of image file paths (one per page),
    or an empty list if conversion is not possible.
    """
    ext = file_path.rsplit('.', 1)[-1].lower() if '.' in file_path else ''
    if ext in IMAGE_EXTS or not HAS_PYMUPDF:
        return []

    pdf_path = file_path
    cleanup_pdf = False

    if ext != 'pdf':
        converted = None
        if ext in OFFICE_EXTS:
            converted = _convert_office_to_pdf(file_path)
        if not converted and ext in PLAIN_TEXT_EXTS:
            converted = _convert_text_to_pdf(file_path)
        if not converted and ext in ('docx', 'doc'):
            converted = _convert_docx_to_pdf_reportlab(file_path)
        if not converted and ext in OFFICE_EXTS:
            # Office conversion failed and it's not a text file — give up
            return []
        if not converted:
            return []
        pdf_path = converted
        cleanup_pdf = True

    try:
        doc = pymupdf.open(pdf_path)
        images = []
        for page in doc:
            pix = page.get_pixmap(dpi=150)
            img_path = os.path.join(
                app.config['UPLOAD_FOLDER'],
                'docpage_%s.jpg' % uuid.uuid4().hex,
            )
            mode = 'RGBA' if pix.alpha else 'RGB'
            pil_page = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
            if pil_page.mode != 'RGB':
                pil_page = pil_page.convert('RGB')
            pil_page.save(img_path, 'JPEG', quality=72)
            images.append(img_path)
        doc.close()
        return images
    except Exception:
        return []
    finally:
        if cleanup_pdf:
            try:
                os.unlink(pdf_path)
            except OSError:
                pass


# ─── PAGE NUMBERING ────────────────────────────────────────────────────────────

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_footer()
            super().showPage()
        super().save()

    def draw_page_footer(self):
        self.setFont("Times-Bold", 14)
        self.setFillColor(colors.HexColor("#333333"))
        width, height = letter
        self.drawRightString(width - 0.5 * inch, 0.35 * inch, "Page %d" % self._pageNumber)
        self.setFont("Times-Roman", 11)
        self.setFillColor(colors.HexColor("#666666"))
        self.drawString(0.5 * inch, 0.35 * inch, "Evidence Brief")
        self.setStrokeColor(colors.HexColor("#cccccc"))
        self.setLineWidth(0.5)
        self.line(0.5 * inch, 0.6 * inch, width - 0.5 * inch, 0.6 * inch)


# ─── TAB DIVIDER FLOWABLE ──────────────────────────────────────────────────────

class TabDivider(Flowable):
    """Blank tab-divider page.
    Records self.canv._pageNumber in page_registry on first draw (TOC two-pass).
    """
    def __init__(self, tab_number, tab_title, width, height,
                 page_registry=None, registry_key=None):
        super().__init__()
        self.tab_number = tab_number
        self.tab_title = tab_title
        self.width = width
        self.height = height
        self.page_registry = page_registry
        self.registry_key = registry_key if registry_key is not None else tab_number

    def draw(self):
        # Record page number for TOC (pass 1 only)
        if self.page_registry is not None:
            self.page_registry[self.registry_key] = self.canv._pageNumber

        cx = self.width / 2
        cy = self.height / 2
        rule_half_w = self.width * 0.35

        # Thin rules flanking the text block
        self.canv.setStrokeColor(colors.black)
        self.canv.setLineWidth(0.5)
        top_y = cy + 0.55 * inch
        bot_y = cy - (0.45 * inch if self.tab_title else 0.15 * inch)
        self.canv.line(cx - rule_half_w, top_y, cx + rule_half_w, top_y)
        self.canv.line(cx - rule_half_w, bot_y, cx + rule_half_w, bot_y)

        # "TAB N"
        self.canv.setFillColor(colors.black)
        self.canv.setFont("Times-Bold", 28)
        label_y = cy + (0.3 * inch if self.tab_title else 0.1 * inch)
        self.canv.drawCentredString(cx, label_y, "TAB %d" % self.tab_number)

        # Tab title
        if self.tab_title:
            self.canv.setFont("Times-Roman", 18)
            self.canv.drawCentredString(cx, cy - 0.25 * inch, self.tab_title)

    def wrap(self, availWidth, availHeight):
        return self.width, self.height


# ─── IMAGE COMPRESSION HELPER ──────────────────────────────────────────────────

def _rl_image_compressed(img_path, max_w_pts, max_h_pts, quality=72):
    """Return a centred RLImage downscaled to 150 DPI at its display size and
    re-encoded as JPEG.  This keeps individual image contributions small without
    visible quality loss at typical screen/print sizes."""
    TARGET_DPI = 150
    PTS_PER_INCH = 72.0
    with Image.open(img_path) as pil:
        orig_w, orig_h = pil.size
        scale = min(max_w_pts / orig_w, max_h_pts / orig_h, 1.0)
        disp_w = orig_w * scale   # points
        disp_h = orig_h * scale
        # Pixel dimensions needed for TARGET_DPI at the display size
        px_w = max(1, int(disp_w / PTS_PER_INCH * TARGET_DPI))
        px_h = max(1, int(disp_h / PTS_PER_INCH * TARGET_DPI))
        img = pil.copy()
    # Downscale only — never upscale
    if img.width > px_w or img.height > px_h:
        img = img.resize((px_w, px_h), Image.LANCZOS)
    if img.mode != 'RGB':
        img = img.convert('RGB')
    buf = io.BytesIO()
    img.save(buf, 'JPEG', quality=quality, optimize=True)
    buf.seek(0)
    rl = RLImage(buf, width=disp_w, height=disp_h)
    rl.hAlign = 'CENTER'
    return rl


# ─── STORY BUILDER ─────────────────────────────────────────────────────────────

def _build_story(case_info, tabs_data, usable_w, usable_h,
                 page_registry=None, tab_pages=None, text_cache=None,
                 doc_page_images=None):
    """
    Build and return the full ReportLab story.

    page_registry   – dict populated with {tab_idx: page_number} during draw (pass 1).
    tab_pages       – dict {tab_idx: page_number} used to fill the TOC (pass 2).
    text_cache      – shared dict {path: (text, truncated, ok)} to avoid double extraction.
    doc_page_images – dict {path: [img_path, …]} of pre-rendered document pages.
    """
    if text_cache is None:
        text_cache = {}

    styles = getSampleStyleSheet()
    story = []

    # ── TITLE PAGE (LTB template style) ──────────────────────────────────────
    # Paragraph styles — Times New Roman throughout, matching the LTB template
    tn_r  = ParagraphStyle('TnR',  parent=styles['Normal'],
                           fontName='Times-Roman',  fontSize=12, leading=18)
    tn_b  = ParagraphStyle('TnB',  parent=styles['Normal'],
                           fontName='Times-Bold',   fontSize=12, leading=18)
    tn_rc = ParagraphStyle('TnRc', parent=styles['Normal'],
                           fontName='Times-Roman',  fontSize=12, leading=18,
                           alignment=TA_CENTER)
    tn_bc = ParagraphStyle('TnBc', parent=styles['Normal'],
                           fontName='Times-Bold',   fontSize=12, leading=18,
                           alignment=TA_CENTER)
    tn_rr = ParagraphStyle('TnRr', parent=styles['Normal'],
                           fontName='Times-Roman',  fontSize=12, leading=18,
                           alignment=TA_RIGHT)

    # File number — top right
    fn = case_info.get('file_number', '')
    if fn:
        story.append(Paragraph("File Number: " + xml_escape(fn), tn_rr))
    story.append(Spacer(1, 0.45 * inch))

    # LANDLORD AND TENANT BOARD / TRIBUNALS ONTARIO
    story.append(Paragraph("LANDLORD AND TENANT BOARD", tn_bc))
    story.append(Paragraph("TRIBUNALS ONTARIO", tn_bc))
    story.append(Spacer(1, 0.35 * inch))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black))
    story.append(Spacer(1, 0.3 * inch))

    # In the matter of (rental unit address = respondent_address)
    unit_addr = case_info.get('respondent_address', '')
    if unit_addr:
        story.append(Paragraph("In the matter of: " + xml_escape(unit_addr), tn_r))
        story.append(Spacer(1, 0.3 * inch))

    # Between
    story.append(Paragraph("Between:", tn_r))
    story.append(Spacer(1, 0.3 * inch))

    # Party rows: name centered, role right-aligned in same row
    pn_s = ParagraphStyle('PnS', parent=styles['Normal'],
                          fontName='Times-Bold',  fontSize=12, alignment=TA_CENTER)
    pr_s = ParagraphStyle('PrS', parent=styles['Normal'],
                          fontName='Times-Roman', fontSize=12, alignment=TA_RIGHT)
    col_l = usable_w * 0.72
    col_r = usable_w * 0.28

    applicant = xml_escape(case_info.get('applicant_name', ''))
    respondent = xml_escape(case_info.get('respondent_name', ''))

    if applicant:
        t = Table([[Paragraph(applicant, pn_s), Paragraph("Applicant", pr_s)]],
                  colWidths=[col_l, col_r])
        t.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                               ('TOPPADDING', (0,0),(-1,-1), 4),
                               ('BOTTOMPADDING', (0,0),(-1,-1), 4)]))
        story.append(t)

    story.append(Paragraph("and", tn_rc))

    if respondent:
        t = Table([[Paragraph(respondent, pn_s), Paragraph("Respondent", pr_s)]],
                  colWidths=[col_l, col_r])
        t.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                               ('TOPPADDING', (0,0),(-1,-1), 4),
                               ('BOTTOMPADDING', (0,0),(-1,-1), 4)]))
        story.append(t)

    story.append(Spacer(1, 0.9 * inch))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black))
    story.append(Spacer(1, 0.18 * inch))
    story.append(Paragraph("TENANT&#8217;S EVIDENCE BRIEF", tn_bc))
    story.append(Spacer(1, 0.18 * inch))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black))
    story.append(PageBreak())

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────────
    story.append(Paragraph("TABLE OF CONTENTS", tn_bc))
    story.append(Spacer(1, 0.1 * inch))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.black))
    story.append(Spacer(1, 0.2 * inch))

    toc_entry_s = ParagraphStyle('TocE', parent=styles['Normal'],
                                 fontName='Times-Roman', fontSize=12, spaceAfter=4)
    toc_hdr_s   = ParagraphStyle('TocH', parent=styles['Normal'],
                                 fontName='Times-Bold',  fontSize=12, spaceAfter=4)
    toc_page_s  = ParagraphStyle('TocP', parent=styles['Normal'],
                                 fontName='Times-Roman', fontSize=12,
                                 spaceAfter=4, alignment=TA_RIGHT)

    toc_data = [[
        Paragraph("Tab",         toc_hdr_s),
        Paragraph("Description", toc_hdr_s),
        Paragraph("Page",        toc_hdr_s),
    ]]
    for i, tab in enumerate(tabs_data, 1):
        pg = str(tab_pages[i]) if (tab_pages and i in tab_pages) else '—'
        toc_data.append([
            Paragraph(str(i),                                toc_entry_s),
            Paragraph(tab.get('title') or ("Tab %d" % i),   toc_entry_s),
            Paragraph(pg,                                    toc_page_s),
        ])

    toc_table = Table(toc_data, colWidths=[0.6 * inch, usable_w - 1.6 * inch, 1 * inch])
    toc_table.setStyle(TableStyle([
        ('TOPPADDING',    (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING',   (0, 0), (-1, -1), 0),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
        ('LINEBELOW',     (0, 0), (-1,  0), 0.5, colors.black),
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN',         (2, 0), (2,  -1), 'RIGHT'),
    ]))
    story.append(toc_table)
    story.append(PageBreak())

    # ── TAB SECTIONS ──────────────────────────────────────────────────────────
    photo_caption_style = ParagraphStyle(
        'PhotoCaption', parent=styles['Normal'],
        fontName='Times-Roman', fontSize=9,
        textColor=colors.black,
        alignment=TA_CENTER, spaceAfter=16,
    )
    no_file_style = ParagraphStyle(
        'NoFile', parent=styles['Normal'],
        fontName='Times-Roman', fontSize=11,
        textColor=colors.grey, alignment=TA_CENTER,
    )
    doc_body_style = ParagraphStyle(
        'DocBody', parent=styles['Normal'],
        fontName='Times-Roman', fontSize=10,
        leading=14, spaceAfter=4,
    )
    doc_note_style = ParagraphStyle(
        'DocNote', parent=styles['Normal'],
        fontName='Times-Italic', fontSize=9,
        textColor=colors.grey, spaceAfter=8,
    )

    for tab_idx, tab in enumerate(tabs_data, 1):
        story.append(TabDivider(
            tab_number=tab_idx,
            tab_title=tab.get('title', ''),
            width=usable_w,
            height=usable_h,
            page_registry=page_registry,
            registry_key=tab_idx,
        ))
        story.append(PageBreak())

        files = tab.get('images', [])
        if not files:
            story.append(Spacer(1, 2 * inch))
            story.append(Paragraph("No files uploaded for this tab.", no_file_style))
            story.append(PageBreak())
            continue

        img_files = [(i, p) for i, p in enumerate(files) if is_image_file(p)]
        doc_files = [(i, p) for i, p in enumerate(files) if not is_image_file(p)]

        # Embed images — one per page
        img_counter = 0
        for orig_idx, img_path in img_files:
            img_counter += 1
            try:
                rl_img = _rl_image_compressed(
                    img_path,
                    max_w_pts=usable_w,
                    max_h_pts=usable_h - 0.6 * inch,
                )
                story.append(rl_img)
                title_part = (" - " + tab.get('title')) if tab.get('title') else ""
                story.append(Paragraph(
                    "Tab %d%s | Photo %d" % (tab_idx, title_part, img_counter),
                    photo_caption_style
                ))
            except Exception:
                story.append(Paragraph(
                    "[Image could not be loaded: %s]" % os.path.basename(img_path),
                    photo_caption_style
                ))
            story.append(PageBreak())

        # Non-image documents: embed converted pages or fall back to text
        for orig_idx, doc_path in doc_files:
            page_imgs = (doc_page_images or {}).get(doc_path, [])

            if page_imgs:
                # ── Converted document: embed each page directly ──
                for pg_idx, pg_img_path in enumerate(page_imgs):
                    try:
                        rl_img = _rl_image_compressed(
                            pg_img_path,
                            max_w_pts=usable_w,
                            max_h_pts=usable_h,
                        )
                        story.append(rl_img)
                    except Exception:
                        story.append(Paragraph(
                            "[Page could not be rendered]",
                            doc_note_style,
                        ))
                    story.append(PageBreak())
            else:
                # ── Fallback: extract and display full text content ──
                if doc_path not in text_cache:
                    text_cache[doc_path] = extract_text_from_file(doc_path)
                text, ok = text_cache[doc_path]

                if ok and text and text.strip():
                    clean = text.replace('\r\n', '\n').replace('\r', '\n')
                    safe  = xml_escape(clean)
                    html  = safe.replace('\n\n', '<br/><br/>').replace('\n', '<br/>')
                    story.append(Paragraph(html, doc_body_style))
                elif not ok:
                    story.append(Paragraph(
                        "Binary or unsupported format — install PyMuPDF and "
                        "LibreOffice for full document conversion.",
                        doc_note_style,
                    ))
                else:
                    story.append(Paragraph(
                        "[File appears to be empty or contains no extractable text.]",
                        doc_note_style,
                    ))
                story.append(PageBreak())

    return story


# ─── PDF GENERATION (two-pass for accurate TOC page numbers) ──────────────────

def generate_evidence_brief(case_info, tabs_data, output_path):
    page_w, page_h = letter
    doc_kwargs = dict(
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    usable_w = page_w - 1.5 * inch
    usable_h = page_h - 1.5 * inch - 0.5 * inch

    # Pre-convert non-image documents to rendered page images
    doc_page_images = {}
    for tab in tabs_data:
        for f in tab.get('images', []):
            if not is_image_file(f) and f not in doc_page_images:
                page_imgs = convert_document_to_images(f)
                if page_imgs:
                    doc_page_images[f] = page_imgs

    # Shared text cache so documents are only read once across both passes
    text_cache = {}

    # Pass 1 — dry run to BytesIO; populates page_registry via TabDivider.draw()
    page_registry = {}
    doc1 = SimpleDocTemplate(io.BytesIO(), **doc_kwargs)
    story1 = _build_story(case_info, tabs_data, usable_w, usable_h,
                          page_registry=page_registry,
                          tab_pages=None,
                          text_cache=text_cache,
                          doc_page_images=doc_page_images)
    doc1.build(story1, canvasmaker=NumberedCanvas)

    # Pass 2 — real render with accurate TOC page numbers
    doc2 = SimpleDocTemplate(output_path, **doc_kwargs)
    story2 = _build_story(case_info, tabs_data, usable_w, usable_h,
                          page_registry=None,
                          tab_pages=page_registry,
                          text_cache=text_cache,
                          doc_page_images=doc_page_images)
    doc2.build(story2, canvasmaker=NumberedCanvas)


# ─── ROUTES ───────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html')


def _cleanup_old_uploads():
    """Remove uploaded files older than 1 hour to prevent disk fill."""
    folder = app.config['UPLOAD_FOLDER']
    cutoff = time.time() - 3600
    try:
        for fname in os.listdir(folder):
            fpath = os.path.join(folder, fname)
            if os.path.isfile(fpath) and os.path.getmtime(fpath) < cutoff:
                try:
                    os.remove(fpath)
                except OSError:
                    pass
    except FileNotFoundError:
        pass


@app.route('/upload', methods=['POST'])
def upload_files():
    _cleanup_old_uploads()
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400

    saved = []
    for f in request.files.getlist('files'):
        if not f or not f.filename:
            continue
        original_name = f.filename
        ext = original_name.rsplit('.', 1)[-1].lower() if '.' in original_name else 'bin'
        fname = "%s.%s" % (uuid.uuid4().hex, ext)
        path = os.path.join(app.config['UPLOAD_FOLDER'], fname)
        f.save(path)
        # Convert HEIC/HEIF to JPEG so PIL and ReportLab can handle it
        if ext in HEIC_EXTS:
            jpeg_name = '%s.jpg' % uuid.uuid4().hex
            jpeg_path = os.path.join(app.config['UPLOAD_FOLDER'], jpeg_name)
            if _convert_heic_to_jpeg(path, jpeg_path):
                try:
                    os.unlink(path)
                except OSError:
                    pass
                path, fname, ext = jpeg_path, jpeg_name, 'jpg'
        thumb_name = None
        is_image = False
        try:
            with Image.open(path) as img:
                img.thumbnail((300, 300))
                thumb_name = "thumb_%s" % fname
                thumb_path = os.path.join(app.config['UPLOAD_FOLDER'], thumb_name)
                img.save(thumb_path)
                is_image = True
        except Exception:
            pass
        saved.append({'id': fname, 'thumb': thumb_name, 'original': fname,
                      'name': original_name, 'is_image': is_image})

    return jsonify({'files': saved})


@app.route('/thumbnail/<filename>')
def get_thumbnail(filename):
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(path):
        return send_file(path)
    return '', 404


# ─── EVIDENCE STATS ────────────────────────────────────────────────────────────

DEFAULT_EVIDENCE_STATS = {
    "N4": {
        "description": "Non-payment of rent",
        "total_cases_analyzed": 0,
        "evidence_types": [
            {"category": "N4 Notice of Termination", "percentage": 92},
            {"category": "Financial records (rent receipts, bank statements, rent ledger)", "percentage": 78},
            {"category": "Lease agreement", "percentage": 65},
            {"category": "Payment history / transaction records", "percentage": 58},
            {"category": "Communication records (emails, text messages, letters)", "percentage": 42},
            {"category": "Legal documents (prior orders, court filings)", "percentage": 28},
            {"category": "Witness testimony", "percentage": 22},
            {"category": "Government/third-party records (inspection reports, municipal notices)", "percentage": 15},
            {"category": "Photos of unit conditions", "percentage": 12},
            {"category": "Maintenance/repair requests or records", "percentage": 8},
        ]
    }
}


@app.route('/evidence-stats')
def evidence_stats():
    stats_path = os.path.join(os.path.dirname(__file__), 'evidence_stats.json')
    if os.path.exists(stats_path):
        try:
            with open(stats_path, 'r') as f:
                return jsonify(json.load(f))
        except (json.JSONDecodeError, IOError):
            pass
    return jsonify(DEFAULT_EVIDENCE_STATS)


@app.route('/generate', methods=['POST'])
def generate():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    case_info = data.get('case_info', {})
    tabs_data = data.get('tabs', [])

    for tab in tabs_data:
        resolved = []
        for img_id in tab.get('images', []):
            path = os.path.join(app.config['UPLOAD_FOLDER'], img_id)
            if os.path.exists(path):
                resolved.append(path)
        tab['images'] = resolved

    output_path = os.path.join(
        app.config['UPLOAD_FOLDER'], "brief_%s.pdf" % uuid.uuid4().hex
    )

    try:
        generate_evidence_brief(case_info, tabs_data, output_path)
        download_name = "Evidence_Brief_%s.pdf" % case_info.get('file_number', 'LTB')
        return send_file(
            output_path,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=download_name,
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    debug = os.environ.get('FLASK_DEBUG', 'false').lower() == 'true'
    port = int(os.environ.get('PORT', 5050))
    app.run(debug=debug, host='0.0.0.0', port=port)
